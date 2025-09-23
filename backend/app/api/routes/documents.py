import os
import uuid
import json
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form, BackgroundTasks
from fastapi.responses import FileResponse
from sqlalchemy.orm import Session
from typing import List, Optional
from pathlib import Path
try:
    import PyPDF2
except ImportError:
    PyPDF2 = None
try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None
from app.db.database import get_db
from app.core.security import get_current_active_user
from app.models.models import User, Document, DocumentChunk, ChunkEmbedding
from app.schemas.schemas import Document as DocumentSchema, DocumentCreate
from app.services.document_service import document_processor
from app.core.config import settings

router = APIRouter()


def process_document_background(document_id: int, file_path: str, file_type: str, db: Session):
    """Background task to process document"""
    try:
        # Process the document
        result = document_processor.process_document(file_path, file_type)
        
        # Update document status
        document = db.query(Document).filter(Document.id == document_id).first()
        
        if result['success']:
            # Save chunks and embeddings
            for chunk_data in result['chunks']:
                # Create chunk
                chunk = DocumentChunk(
                    document_id=document_id,
                    content=chunk_data['content'],
                    chunk_index=chunk_data['chunk_index'],
                    doc_metadata=chunk_data['metadata']
                )
                db.add(chunk)
                db.flush()  # Get the chunk ID
                
                # Create embedding
                embedding = ChunkEmbedding(
                    chunk_id=chunk.id,
                    embedding=json.dumps(chunk_data['embedding'])
                )
                db.add(embedding)
            
            document.status = "completed"
        else:
            document.status = "failed"
        
        db.commit()
        
    except Exception as e:
        # Update document status to failed
        document = db.query(Document).filter(Document.id == document_id).first()
        if document:
            document.status = "failed"
            db.commit()


@router.post("/upload", response_model=DocumentSchema)
async def upload_document(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    title: str = Form(...),
    description: Optional[str] = Form(None),
    current_user: User = Depends(get_current_active_user),
    db: Session = Depends(get_db)
):
    """Upload a new document"""
    
    # Validate file type
    if file.content_type not in ["application/pdf", "text/plain", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"]:
        raise HTTPException(status_code=400, detail="Unsupported file type")
    
    # Validate file size
    if file.size > settings.MAX_FILE_SIZE:
        raise HTTPException(status_code=400, detail="File too large")
    
    # Generate unique filename
    file_extension = Path(file.filename).suffix.lower()
    unique_filename = f"{uuid.uuid4()}{file_extension}"
    file_path = os.path.join(settings.UPLOAD_DIR, unique_filename)
    
    # Ensure upload directory exists
    os.makedirs(settings.UPLOAD_DIR, exist_ok=True)
    
    # Save file
    with open(file_path, "wb") as buffer:
        content = await file.read()
        buffer.write(content)
    
    # Determine file type
    file_type_mapping = {
        ".pdf": "pdf",
        ".txt": "txt",
        ".md": "md",
        ".docx": "docx"
    }
    file_type = file_type_mapping.get(file_extension, "unknown")
    
    # Create document record
    document = Document(
        user_id=current_user.id,
        title=title,
        description=description,
        filename=file.filename,
        file_path=file_path,
        file_type=file_type,
        file_size=file.size,
        status="processing"
    )
    
    db.add(document)
    db.commit()
    db.refresh(document)
    
    # Process document in background
    background_tasks.add_task(process_document_background, document.id, file_path, file_type, db)
    
    return document


@router.get("/", response_model=List[DocumentSchema])
async def get_documents(
    skip: int = 0,
    limit: int = 20,
    current_user: User = Depends(get_current_active_user),
    db: Session = Depends(get_db)
):
    """Get user's documents"""
    documents = db.query(Document).filter(
        Document.user_id == current_user.id
    ).offset(skip).limit(limit).all()
    
    return documents


@router.get("/{document_id}", response_model=DocumentSchema)
async def get_document(
    document_id: int,
    current_user: User = Depends(get_current_active_user),
    db: Session = Depends(get_db)
):
    """Get a specific document"""
    document = db.query(Document).filter(
        Document.id == document_id,
        Document.user_id == current_user.id
    ).first()
    
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    
    return document


@router.delete("/{document_id}")
async def delete_document(
    document_id: int,
    current_user: User = Depends(get_current_active_user),
    db: Session = Depends(get_db)
):
    """Delete a document"""
    document = db.query(Document).filter(
        Document.id == document_id,
        Document.user_id == current_user.id
    ).first()
    
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    
    # Delete file from filesystem
    try:
        if os.path.exists(document.file_path):
            os.remove(document.file_path)
    except Exception:
        pass  # Continue even if file deletion fails
    
    # Delete from database (cascades to chunks and embeddings)
    db.delete(document)
    db.commit()
    
    return {"message": "Document deleted successfully"}


@router.get("/status", response_model=dict)
async def get_processing_status(
    current_user: User = Depends(get_current_active_user),
    db: Session = Depends(get_db)
):
    """Get processing status of all user documents"""
    documents = db.query(Document).filter(
        Document.user_id == current_user.id
    ).all()
    
    status_summary = {
        "total": len(documents),
        "completed": len([d for d in documents if d.status == "completed"]),
        "processing": len([d for d in documents if d.status == "processing"]),
        "failed": len([d for d in documents if d.status == "failed"]),
        "documents": [
            {
                "id": doc.id,
                "title": doc.title,
                "status": doc.status,
                "updated_at": doc.updated_at
            }
            for doc in documents
        ]
    }
    
    return status_summary


@router.get("/{document_id}/content")
async def get_document_content(
    document_id: int,
    current_user: User = Depends(get_current_active_user),
    db: Session = Depends(get_db)
):
    """Get full document content for preview"""
    document = db.query(Document).filter(
        Document.id == document_id,
        Document.user_id == current_user.id
    ).first()
    
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    
    if not os.path.exists(document.file_path):
        raise HTTPException(status_code=404, detail="File not found")
    
    try:
        # Read the full document content based on file type
        if document.file_type == "txt" or document.file_type == "md":
            with open(document.file_path, 'r', encoding='utf-8') as file:
                content = file.read()
        elif document.file_type == "pdf":
            # For PDF, we'll extract text content
            import PyPDF2
            content = ""
            with open(document.file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    content += page.extract_text() + "\n\n"
        elif document.file_type == "docx":
            # For DOCX, extract text content
            from docx import Document as DocxDocument
            docx_doc = DocxDocument(document.file_path)
            content = ""
            for paragraph in docx_doc.paragraphs:
                content += paragraph.text + "\n"
        else:
            content = "Preview not available for this file type."
            
    except Exception as e:
        content = f"Error reading document: {str(e)}"
    
    return {
        "document": document,
        "content": content
    }


@router.get("/{document_id}/download")
async def download_document(
    document_id: int,
    current_user: User = Depends(get_current_active_user),
    db: Session = Depends(get_db)
):
    """Download a document"""
    document = db.query(Document).filter(
        Document.id == document_id,
        Document.user_id == current_user.id
    ).first()
    
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    
    if not os.path.exists(document.file_path):
        raise HTTPException(status_code=404, detail="File not found")
    
    return FileResponse(
        document.file_path,
        filename=document.filename,
        media_type="application/octet-stream"
    )